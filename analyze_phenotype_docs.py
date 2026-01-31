"""
Pass 1: Document Analysis for PheKB Phenotypes

This script analyzes PDF and Word documents from phenotype directories,
extracting clinical codes and criteria using Claude.

Output: data/phekb-raw/{phenotype-id}/document_analysis.json
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from datetime import datetime
import logging
import re

# Add backend to path
backend_path = Path(__file__).parent / "backend"
sys.path.insert(0, str(backend_path))

from pypdf import PdfReader
from docx import Document
from anthropic import Anthropic
from dotenv import load_dotenv

# Load environment
load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(str(pdf_path))
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n\n".join(text)
    except Exception as e:
        logger.warning(f"Could not extract text from PDF {pdf_path.name}: {e}")
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from Word document (.docx)"""
    try:
        doc = Document(str(docx_path))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n\n".join(text)
    except Exception as e:
        logger.warning(f"Could not extract text from Word doc {docx_path.name}: {e}")
        return ""


def extract_text_from_doc(doc_path: Path) -> str:
    """Extract text from old Word document (.doc) using antiword or fallback"""
    try:
        # Try antiword first (common on Linux/WSL)
        result = subprocess.run(
            ["antiword", str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        # Try catdoc as fallback
        result = subprocess.run(
            ["catdoc", str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Try python-docx anyway (sometimes works on .doc files)
    try:
        doc = Document(str(doc_path))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n\n".join(text)
    except Exception:
        pass

    logger.warning(f"Could not extract text from .doc file {doc_path.name} - install antiword or catdoc")
    return ""


def extract_all_documents(phenotype_dir: Path) -> dict:
    """Extract text from all documents in a phenotype directory"""
    file_contents = {}

    for file_path in phenotype_dir.iterdir():
        if file_path.suffix.lower() == ".pdf":
            logger.info(f"    Extracting: {file_path.name}")
            text = extract_text_from_pdf(file_path)
            if text:
                file_contents[file_path.name] = text
        elif file_path.suffix.lower() == ".docx":
            logger.info(f"    Extracting: {file_path.name}")
            text = extract_text_from_docx(file_path)
            if text:
                file_contents[file_path.name] = text
        elif file_path.suffix.lower() == ".doc":
            logger.info(f"    Extracting: {file_path.name}")
            text = extract_text_from_doc(file_path)
            if text:
                file_contents[file_path.name] = text

    return file_contents


def analyze_documents_with_claude(
    phenotype_id: str,
    file_contents: dict,
    description: str,
    client: Anthropic
) -> dict:
    """Use Claude to extract structured data from documents"""

    system_prompt = """You are a clinical informatics expert analyzing phenotype definition documents.

Your task is to extract ALL clinical codes and criteria found in these documents.

IMPORTANT:
- Extract ONLY codes that are explicitly mentioned in the documents
- Include the exact context where each code was found
- Identify clinical criteria and thresholds (lab values, etc.)
- Summarize the phenotype algorithm logic

Code systems to look for:
- ICD-9-CM: Format like 250.00, 401.1, etc.
- ICD-10-CM: Format like E11, I10, J45.20, etc.
- LOINC: Format like 4548-4, 2339-0, etc.
- SNOMED CT: Numeric codes, often 6+ digits
- RxNorm: Medication codes
- CPT: Procedure codes like 99213, 93000

Be thorough - examine all document content carefully for codes."""

    # Build document content for analysis
    doc_parts = []
    for filename, content in file_contents.items():
        # Include more content per file for thorough analysis
        truncated = content[:8000]
        if len(content) > 8000:
            truncated += "\n... (document continues)"
        doc_parts.append(f"=== {filename} ===\n{truncated}")

    documents_text = "\n\n".join(doc_parts)

    user_prompt = f"""Analyze these phenotype documents for: {phenotype_id}

Description: {description[:1000]}

DOCUMENTS:
{documents_text}

Extract and return as JSON:
{{
  "extracted_codes": [
    {{
      "system": "ICD-10-CM",  // One of: ICD-9-CM, ICD-10-CM, LOINC, SNOMED CT, RxNorm, CPT
      "code": "E11",
      "display": "Type 2 diabetes mellitus",
      "found_in": "filename.pdf",
      "context": "Quote or description of where this code appears"
    }}
  ],
  "clinical_criteria": [
    "HbA1c >= 6.5%",
    "Fasting glucose >= 126 mg/dL"
  ],
  "algorithm_summary": "Brief description of how the phenotype algorithm works"
}}

Important:
- ONLY include codes explicitly found in the documents
- Include ALL codes found, even if there are many
- Provide the filename where each code was found
- Include context/quotes for each code"""

    try:
        logger.info("    Calling Claude API for document analysis...")
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16384,  # Increased for complex phenotypes
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        response_text = message.content[0].text

        # Check for truncation
        if message.stop_reason == "max_tokens":
            logger.warning("    Response was truncated (max_tokens reached)")

        # Extract JSON from response
        json_text = None

        if "```json" in response_text:
            json_start = response_text.index("```json") + 7
            # Find closing fence AFTER the opening one
            remaining = response_text[json_start:]
            if "```" in remaining:
                json_end = json_start + remaining.index("```")
                json_text = response_text[json_start:json_end].strip()
            else:
                # No closing fence - response was truncated
                # Try to salvage by finding the last complete object
                json_text = remaining.strip()
                logger.warning("    No closing fence found, attempting to salvage JSON")
        elif "```" in response_text:
            json_start = response_text.index("```") + 3
            remaining = response_text[json_start:]
            if "```" in remaining:
                json_end = json_start + remaining.index("```")
                json_text = response_text[json_start:json_end].strip()
            else:
                json_text = remaining.strip()
        else:
            json_text = response_text.strip()

        # Try to fix truncated JSON
        if json_text:
            # Remove any leading non-JSON text
            if not json_text.startswith("{"):
                brace_pos = json_text.find("{")
                if brace_pos != -1:
                    json_text = json_text[brace_pos:]

            # Try to parse, and if it fails due to truncation, try to repair
            try:
                result = json.loads(json_text)
            except json.JSONDecodeError:
                # Try to repair truncated JSON by closing open structures
                repaired = repair_truncated_json(json_text)
                if repaired:
                    result = json.loads(repaired)
                else:
                    raise
        else:
            raise ValueError("No JSON content found in response")

        logger.info(f"    Found {len(result.get('extracted_codes', []))} codes")
        return result

    except Exception as e:
        logger.error(f"    Failed to analyze with Claude: {e}")
        raise


def repair_truncated_json(json_text: str) -> str:
    """Attempt to repair truncated JSON by closing open structures"""
    # Count open braces and brackets
    open_braces = json_text.count("{") - json_text.count("}")
    open_brackets = json_text.count("[") - json_text.count("]")

    if open_braces < 0 or open_brackets < 0:
        return None  # Malformed, not just truncated

    # Find the last complete array element or object
    # Look for the last complete "}" that could end an object in an array
    last_complete = json_text.rfind("},")
    if last_complete == -1:
        last_complete = json_text.rfind("}")

    if last_complete > 0:
        # Truncate at the last complete object
        json_text = json_text[:last_complete + 1]

        # Recount after truncation
        open_braces = json_text.count("{") - json_text.count("}")
        open_brackets = json_text.count("[") - json_text.count("]")

        # Close remaining structures
        json_text += "]" * open_brackets + "}" * open_braces

        return json_text

    return None


def save_document_analysis(
    phenotype_dir: Path,
    phenotype_id: str,
    analyzed_files: list,
    analysis_result: dict
):
    """Save the document analysis to JSON"""

    output = {
        "phenotype_id": phenotype_id,
        "analyzed_files": analyzed_files,
        "extracted_codes": analysis_result.get("extracted_codes", []),
        "clinical_criteria": analysis_result.get("clinical_criteria", []),
        "algorithm_summary": analysis_result.get("algorithm_summary", ""),
        "analysis_timestamp": datetime.utcnow().isoformat() + "Z"
    }

    output_file = phenotype_dir / "document_analysis.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2)

    logger.info(f"    Saved: {output_file}")


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Analyze phenotype documents")
    parser.add_argument("--limit", type=int, help="Limit number of phenotypes to process")
    parser.add_argument("--force", action="store_true", help="Re-analyze even if document_analysis.json exists")
    parser.add_argument("--phenotype", type=str, help="Process specific phenotype ID only")
    args = parser.parse_args()

    # Check for API key
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        logger.error("ANTHROPIC_API_KEY not found in environment")
        logger.error("Add it to your .env file")
        sys.exit(1)

    # Initialize Claude client
    client = Anthropic(api_key=api_key)

    # Setup directories
    data_dir = Path("data/phekb-raw")

    # Get phenotype directories
    if args.phenotype:
        phenotype_dirs = [data_dir / args.phenotype]
        if not phenotype_dirs[0].exists():
            logger.error(f"Phenotype directory not found: {args.phenotype}")
            sys.exit(1)
    else:
        phenotype_dirs = [
            d for d in data_dir.iterdir()
            if d.is_dir() and d.name != "__pycache__"
        ]

    if args.limit:
        phenotype_dirs = phenotype_dirs[:args.limit]

    logger.info("=" * 60)
    logger.info(f"Document Analysis - Pass 1")
    logger.info(f"Processing {len(phenotype_dirs)} phenotypes")
    logger.info("=" * 60)

    success_count = 0
    skipped_count = 0
    no_docs_count = 0
    error_count = 0

    for i, phenotype_dir in enumerate(phenotype_dirs, 1):
        phenotype_id = phenotype_dir.name
        logger.info(f"\n[{i}/{len(phenotype_dirs)}] {phenotype_id}")

        # Check if already analyzed
        analysis_file = phenotype_dir / "document_analysis.json"
        if analysis_file.exists() and not args.force:
            logger.info("  Skipped - already analyzed (use --force to re-analyze)")
            skipped_count += 1
            continue

        try:
            # Extract text from all documents
            logger.info("  Extracting document text...")
            file_contents = extract_all_documents(phenotype_dir)

            if not file_contents:
                logger.info("  No documents found to analyze")
                no_docs_count += 1
                continue

            # Load description
            desc_file = phenotype_dir / "description.txt"
            description = ""
            if desc_file.exists():
                with open(desc_file, "r", encoding="utf-8") as f:
                    description = f.read()

            # Analyze with Claude
            analysis_result = analyze_documents_with_claude(
                phenotype_id,
                file_contents,
                description,
                client
            )

            # Save results
            save_document_analysis(
                phenotype_dir,
                phenotype_id,
                list(file_contents.keys()),
                analysis_result
            )

            success_count += 1

        except Exception as e:
            logger.error(f"  Error: {e}")
            error_count += 1
            continue

    logger.info("\n" + "=" * 60)
    logger.info("Document Analysis Complete!")
    logger.info("=" * 60)
    logger.info(f"Analyzed:    {success_count}")
    logger.info(f"Skipped:     {skipped_count}")
    logger.info(f"No docs:     {no_docs_count}")
    logger.info(f"Errors:      {error_count}")


if __name__ == "__main__":
    main()
